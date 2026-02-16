from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from core.models import SummaryRequest, SummaryResponse
from core.nlp_engine import (
    content_ranking, 
    abstractive_refine, 
    detect_src_language, 
    translate_content, 
    clean_visual_artifacts, 
    extract_and_remove_references, 
    extract_ner
)
import uvicorn
import io
import fitz # PyMuPDF
import docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import re
import asyncio
from concurrent.futures import ThreadPoolExecutor

app = FastAPI(title="Lingual API", version="1.0.0")

# CORS for Frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global Executor for CPU tasks
executor = ThreadPoolExecutor(max_workers=3)

@app.get("/")
def health_check():
    return {"status": "ok", "system": "Multilingual Summarizer Standard"}

def extract_text_from_file(file: UploadFile) -> str:
    content = ""
    try:
        if file.filename.endswith(".pdf"):
            try:
                # Read file bytes into memory
                file_bytes = file.file.read()
                # Open PDF with PyMuPDF (fitz)
                with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                    for page in doc:
                        # get_text() handles multi-column layouts by default (reading order)
                        content += page.get_text() + "\n"
            except Exception as e:
                print(f"PyMuPDF error: {e}")
                if not content:
                     raise ValueError("Could not extract text from PDF. It might be image-based or encrypted.")

        elif file.filename.endswith(".docx"):
            doc = docx.Document(file.file)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"

        elif file.filename.endswith(".txt"):
            # Try utf-8 first, then fallback
            content_bytes = file.file.read()
            try:
                content = content_bytes.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    content = content_bytes.decode("latin-1")
                except:
                    content = content_bytes.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF, DOCX, or TXT.")
            
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"File parsing error: {str(e)}")
    
    # Remove null bytes which can choke some NLP parsers
    return content.replace('\x00', '').strip()

async def run_in_threadpool(func, *args):
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(executor, func, *args)

def execute_summary_logic(text: str, sentences_count: int, language: str, mode: str, ratio: float = None):
    # 0. Global Cleanup & Pre-Processing
    # Clean visual artifacts first (single newlines)
    text = clean_visual_artifacts(text)
    
    # Extract & Remove References/Bibliography to avoid "Citation Trap"
    # This gives us a cleaner text for summarization AND a list of citations
    clean_text, citations_list = extract_and_remove_references(text)
    
    # Extract NER Data from the CLEAN text (or original? usually clean is better for entity context)
    ner_list = extract_ner(clean_text)
    
    # Use the REF-CLEANED text for the rest of the pipeline
    working_text = clean_text

    # 1. Detect Language
    src_lang = detect_src_language(working_text)
    
    # 2. Routing Logic
    summary_sents = []
    ranked_data = [] 
    metrics = {}
    final_summary_text = ""
    
    # Logic Branching
    if mode == "extractive":
        # Extractive: Summarize Source -> Translate Result
        
        summary_sents, ranked_data, metrics = content_ranking(working_text, top_n=sentences_count, ratio=ratio)
        extractive_text = " ".join(summary_sents)
        
        if src_lang != language and language != "auto":
            final_summary_text = translate_content(extractive_text, language)
        else:
            final_summary_text = extractive_text
            
    elif mode == "abstractive":
        # Abstractive: Translate Input -> Summarize Target
        if src_lang != language and language != "auto":
            trans_input = translate_content(working_text, language)
        else:
            trans_input = working_text
            
        # For heatmap, we run ranking on ORIGINAL (cleaned) text to map back sentences.
        _, ranked_data, metrics = content_ranking(working_text, top_n=sentences_count, ratio=ratio)
        
        # Now generate abstractive summary
        final_summary_text = abstractive_refine(trans_input, max_length=150)
        
        summary_sents = [final_summary_text]

    elif mode == "hybrid":
        # Hybrid: Extractive Compression -> Refinement
        
        summary_sents, ranked_data, metrics = content_ranking(working_text, top_n=sentences_count, ratio=ratio)
        extractive_text = " ".join(summary_sents)
        
        if src_lang != language and language != "auto":
            translated_extractive = translate_content(extractive_text, language)
            trans_input = translated_extractive
        else:
            trans_input = extractive_text
            
        final_summary_text = abstractive_refine(trans_input, max_length=150)

    # Fallback if empty
    if not final_summary_text:
        final_summary_text = "Could not generate summary."

    return SummaryResponse(
        original_text=text, # Keep full original text for reference? Yes.
        summary_text=final_summary_text,
        extractive_summary=summary_sents,
        refined_summary=final_summary_text if mode != "extractive" else None,
        ranking_data=ranked_data,
        metrics=metrics,
        citations=citations_list,
        ner_data=ner_list
    )

@app.post("/summarize/text", response_model=SummaryResponse)
async def summarize_text_endpoint(request: SummaryRequest):
    try:
        return await asyncio.wait_for(
            run_in_threadpool(
                execute_summary_logic, 
                request.text, 
                request.sentences_count, 
                request.language, 
                request.mode,
                request.summary_ratio
            ),
            timeout=120.0
        )
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail="Processing timeout (120s limit exceeded)")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/summarize/file", response_model=SummaryResponse)
async def summarize_file_endpoint(
    file: UploadFile = File(...),
    sentences_count: int = Form(5),
    summary_ratio: float = Form(None),
    language: str = Form("en"),
    mode: str = Form("extractive")
):
    try:
        text = extract_text_from_file(file)
    except Exception as e:
         raise HTTPException(status_code=400, detail=f"File parsing error: {e}")

    if not text:
        raise HTTPException(status_code=400, detail="Could not extract text from file.")
        
    if len(text) > 10 * 1024 * 1024:
         raise HTTPException(status_code=413, detail="File too large (Max 10MB text content)")

    try:
        return await asyncio.wait_for(
            run_in_threadpool(
                execute_summary_logic, 
                text, 
                sentences_count, 
                language, 
                mode,
                summary_ratio
            ),
            timeout=180.0
        )
    except asyncio.TimeoutError:
        raise HTTPException(status_code=504, detail="Processing timeout (180s limit exceeded)")
    except Exception as e:
         raise HTTPException(status_code=500, detail=str(e))

@app.post("/export/word")
async def export_word_endpoint(request: SummaryResponse):
    # Generate Docx
    doc = Document()
    doc.add_heading('Executive Summary Report', 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(request.summary_text)
    
    # Add Citations if present
    if request.citations:
        doc.add_heading('Extracted Citations', level=2)
        for i, cit in enumerate(request.citations[:20]): # Limit to first 20 for doc
             doc.add_paragraph(f"[{i+1}] {cit}")
        if len(request.citations) > 20:
            doc.add_paragraph(f"...and {len(request.citations)-20} more.")

    # Add NER Highlights
    if request.ner_data:
        doc.add_heading('Key Entities', level=2)
        # Group by type?
        # Just simple list for now
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Relationship'
        
        for item in request.ner_data[:20]:
            row_cells = table.add_row().cells
            row_cells[0].text = item.name
            row_cells[1].text = item.entity
            row_cells[2].text = item.relationship

    doc.add_heading('Metrics', level=2)
    metrics = request.metrics
    if metrics:
        doc.add_paragraph(f"Original Length: {metrics.get('original_length', 0)} chars")
        doc.add_paragraph(f"Compression Ratio: {metrics.get('compression_ratio', 0)}")
        doc.add_paragraph(f"Sentences: {metrics.get('sentence_count', 0)}")
    
    doc.add_heading('Original Text', level=1)
    
    # Sentence Tokenization & Highlighting
    original_text = request.original_text
    # Truncate if excessively long to safe processing time/file size
    if len(original_text) > 50000:
        original_text = original_text[:50000]
        
    # Prepare comparison set
    summary_sentences = set()
    if request.extractive_summary:
        summary_sentences = {s.strip() for s in request.extractive_summary}
        
    # Split sentences (Regex split looking for sentence terminators followed by space)
    # Using lookbehind to keep the delimiter
    raw_sentences = re.split(r'(?<=[.!?])\s+', original_text)
    
    para = doc.add_paragraph()
    
    for sent in raw_sentences:
        cleaned_sent = sent.strip()
        if not cleaned_sent:
            continue
            
        # Add space restored by split consumption
        run = para.add_run(sent + " ")
        
        if cleaned_sent in summary_sentences:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            run.font.highlight_color = WD_COLOR_INDEX.RED
    
    if len(request.original_text) > 50000:
        doc.add_paragraph("...[Text Truncated for Report]...")
    
    # Save to IO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    headers = {
        'Content-Disposition': 'attachment; filename="summary_report.docx"'
    }
    return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
